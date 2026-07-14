from __future__ import annotations

import logging
import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from PIL import Image, UnidentifiedImageError

from rag_ingestion.config import Settings
from rag_ingestion.models import (
    ExtractedDocument,
    ExtractedImage,
    ExtractedPage,
    ExtractedTable,
    SourceMetadata,
)
from rag_ingestion.utils import compact_whitespace, sha256_bytes, slugify

logger = logging.getLogger(__name__)


class DocxExtractor:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def extract(self, docx_path: Path, file_hash: str, metadata: SourceMetadata) -> ExtractedDocument:
        logger.info("Extracting DOCX: %s", docx_path.name)
        docx = DocxDocument(docx_path)
        text = self._extract_text(docx)
        tables = self._extract_tables(docx)
        images = self._extract_images(
            docx=docx,
            source_docx=docx_path,
            document_text=text,
            metadata=metadata,
        )

        return ExtractedDocument(
            source_file=docx_path.name,
            source_file_stem=docx_path.stem,
            source_path=str(docx_path),
            file_hash=file_hash,
            metadata=metadata,
            pages=[
                ExtractedPage(
                    page_number=1,
                    text=text,
                    tables=tables,
                    images=images,
                )
            ],
        )

    def _extract_text(self, docx) -> str:
        paragraphs = [compact_whitespace(paragraph.text) for paragraph in docx.paragraphs]
        return "\n\n".join(paragraph for paragraph in paragraphs if paragraph).strip()

    def _extract_tables(self, docx) -> list[ExtractedTable]:
        extracted_tables: list[ExtractedTable] = []
        for table_index, table in enumerate(docx.tables, start=1):
            rows = self._normalize_table_rows(
                [[compact_whitespace(cell.text) for cell in row.cells] for row in table.rows]
            )
            if not rows:
                continue
            extracted_tables.append(
                ExtractedTable(
                    page_number=1,
                    table_number=table_index,
                    rows=rows,
                    markdown=self._table_to_markdown(rows),
                )
            )
        return extracted_tables

    def _extract_images(
        self,
        *,
        docx,
        source_docx: Path,
        document_text: str,
        metadata: SourceMetadata,
    ) -> list[ExtractedImage]:
        extracted_images: list[ExtractedImage] = []
        seen_hashes: set[str] = set()

        for related_part in docx.part.related_parts.values():
            content_type = getattr(related_part, "content_type", "")
            if not content_type.startswith("image/"):
                continue

            image_bytes = related_part.blob
            image_hash = sha256_bytes(image_bytes)
            if image_hash in seen_hashes:
                continue
            seen_hashes.add(image_hash)

            image_number = len(extracted_images) + 1
            try:
                width, height = self._save_image_as_png(image_bytes, source_docx, image_number)
            except UnidentifiedImageError:
                logger.warning("Skipping unsupported DOCX image in %s", source_docx.name)
                continue
            except Exception:
                logger.exception("Failed to extract DOCX image %s from %s", image_number, source_docx.name)
                continue

            if width < self.settings.min_image_width or height < self.settings.min_image_height:
                continue

            local_path = self._local_image_path(source_docx, image_number)
            extracted_images.append(
                ExtractedImage(
                    page_number=1,
                    image_number=image_number,
                    width=width,
                    height=height,
                    local_path=str(local_path),
                    storage_path=self._storage_image_path(metadata.vendor, source_docx.stem, image_number),
                    caption=self._find_caption(document_text, image_number),
                    context_text=self._image_context_text(document_text),
                )
            )

        return extracted_images

    def _save_image_as_png(self, image_bytes: bytes, source_docx: Path, image_number: int) -> tuple[int, int]:
        local_path = self._local_image_path(source_docx, image_number)
        local_path.parent.mkdir(parents=True, exist_ok=True)

        with Image.open(BytesIO(image_bytes)) as image:
            width, height = image.size
            if image.mode not in {"RGB", "RGBA"}:
                image = image.convert("RGBA" if "A" in image.getbands() else "RGB")
            image.save(local_path, format="PNG")
            return width, height

    def _local_image_path(self, source_docx: Path, image_number: int) -> Path:
        source_stem = slugify(source_docx.stem, fallback="document")
        return self.settings.extracted_images_dir / source_stem / f"page_1_img_{image_number}.png"

    def _storage_image_path(self, vendor: str, source_file_stem: str, image_number: int) -> str:
        return (
            f"{slugify(vendor, fallback='unknown_vendor')}/"
            f"{slugify(source_file_stem, fallback='document')}/"
            f"page_1_img_{image_number}.png"
        )

    def _normalize_table_rows(self, rows: list[list[str]]) -> list[list[str]]:
        normalized: list[list[str]] = []
        for row in rows:
            normalized_row = [compact_whitespace(cell) for cell in row]
            if any(normalized_row):
                normalized.append(normalized_row)
        return normalized

    def _table_to_markdown(self, rows: list[list[str]]) -> str:
        max_columns = max(len(row) for row in rows)
        padded_rows = [row + [""] * (max_columns - len(row)) for row in rows]
        header = [
            self._escape_markdown_cell(cell) if cell else f"Column {index + 1}"
            for index, cell in enumerate(padded_rows[0])
        ]
        body = padded_rows[1:] or [[""] * max_columns]

        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * max_columns) + " |",
        ]
        for row in body:
            lines.append("| " + " | ".join(self._escape_markdown_cell(cell) for cell in row) + " |")
        return "\n".join(lines)

    def _escape_markdown_cell(self, value: str) -> str:
        return value.replace("|", "\\|").replace("\n", " ").strip()

    def _find_caption(self, text: str, image_number: int) -> str | None:
        caption_pattern = re.compile(
            r"(?im)^\s*((figure|fig\.|image|diagram)\s*"
            + re.escape(str(image_number))
            + r"[\w.\-:)]*\s+.+)$"
        )
        match = caption_pattern.search(text)
        if match:
            return compact_whitespace(match.group(1))
        return None

    def _image_context_text(self, document_text: str, max_chars: int = 2500) -> str:
        text = compact_whitespace(document_text)
        if len(text) <= max_chars:
            return text
        return text[:max_chars].rsplit(" ", 1)[0]
